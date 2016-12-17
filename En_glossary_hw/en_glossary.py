import urllib.request
from docx import Document
from bs4 import BeautifulSoup
import re
import random

def grab_info(term):
	url = "http://pc.net/glossary/definition/" + term
	try:
		html = urllib.request.urlopen(url)
	except urllib.error.HTTPError:
		return None
	html = BeautifulSoup(urllib.request.urlopen(url).read(), 'html.parser')
	word = html.find('h2').getText()

	try:
		definition = html.find('div',{'id':'definition'}).getText()
		definition = re.match(r'(?:[^.:;]+[.:;]){4}', definition).group()
	except AttributeError:
		definition = html.find('div',{'id':'definition'}).getText()

	return (word,definition)

def generate_url(alpha):
	counter = 0
	url = "http://pc.net/glossary/browse/" + alpha
	html = BeautifulSoup(urllib.request.urlopen(url).read(), 'html.parser')
	table = html.find("table", attrs={"class":"list"}).findAll('td', {'class' : 'term'})
	for el in table:
		table[counter] = table[counter].getText().lower()
		table[counter] = ''.join(e for e in table[counter] if e.isalnum())
		counter += 1
	return table

words = []
i = ord('a')
print("working...")
#getting all words
while i <= ord('z'):
	words.extend(generate_url(chr(i)))
	i+=1
random.shuffle(words)

i = 0
b = 0;
#add definitions to doc
document = Document()
while i < 60:
	#print(i+1)
	#print(words[i])
	result = grab_info(words[i])
	if result != None: 
		document.add_heading(str(b+1) + ". " + result[0])
		p = document.add_paragraph(result[1])
		#print(result[0],result[1])
		#print("\n")
		b+=1
	if b == 50:
		break
	i+=1

if b < 50:
	print("Less than 50 definitions!!!")
	print(b)

document.save('demo.docx')
print("Done!")
